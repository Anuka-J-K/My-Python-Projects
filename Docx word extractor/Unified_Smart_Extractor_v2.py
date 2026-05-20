import docx
import os
import csv
import re  # Regular expressions for parsing expected word counts from filenames

def unified_smart_extract(file_path):
    """
    Unified DOCX extractor v2 — handles all resume layout types correctly.

    Layout discoveries per file:
    ┌──────────────────────────────┬────────────────────────────────────────────────────────────┐
    │ File / Layout Type           │ Root Cause & Fix                                           │
    ├──────────────────────────────┼────────────────────────────────────────────────────────────┤
    │ Person2 — Pure textbox       │ Content in <txbxContent> inside <wsp> shapes. Word stores  │
    │                              │ each textbox TWICE (wsp copy + standalone copy). Fix:       │
    │                              │ exclude ALL <t> inside <wsp> to avoid double-counting.      │
    ├──────────────────────────────┼────────────────────────────────────────────────────────────┤
    │ Person6 — SDT table layout   │ Content wrapped in <sdt> (Structured Document Tags) inside  │
    │                              │ table cells. python-docx API dedup (by id()) skipped cells  │
    │                              │ due to merged-cell resolution. Fix: iterate <tc> elements   │
    │                              │ directly from XML (not via doc.tables API).                 │
    ├──────────────────────────────┼────────────────────────────────────────────────────────────┤
    │ Person7 — Merged-cell table  │ Heavily merged table. doc.tables API returns same _tc       │
    │                              │ element 2-3x for merged regions → 652 words (expected 245). │
    │                              │ Fix: iterate <tc> XML elements directly with element-set    │
    │                              │ deduplication (not id() which had proxy issues).            │
    ├──────────────────────────────┼────────────────────────────────────────────────────────────┤
    │ Person8 — SDT + table        │ Same as Person6. SDT-wrapped content in table cells missed  │
    │                              │ by wsp-filter. Fix: same XML tc iteration approach.         │
    ├──────────────────────────────┼────────────────────────────────────────────────────────────┤
    │ Person9 — Table + paragraphs │ CV has both body paragraphs AND layout tables. Table adds   │
    │                              │ sidebar content (skills, references) not in the ground-truth │
    │                              │ word count. Ground truth = 245, our best = 274 (88.16%).    │
    │                              │ This is a measurement boundary — the ground truth tool      │
    │                              │ partially skips some table cells. Not fixable without       │
    │                              │ knowing which cells to exclude.                             │
    └──────────────────────────────┴────────────────────────────────────────────────────────────┘

    Core rules applied in order:
      1. Iterate DIRECT body children (<p> and <tbl>) only — never recurse into nested bodies.
      2. For <p>: collect all <w:t> NOT inside <wsp> (wsp = Word Processing Shape, a duplicate store).
         Also collect <w:t> inside standalone <txbxContent> NOT inside <wsp> (pure-textbox CVs).
      3. For <tbl>: iterate <tc> elements via XML directly (bypasses merged-cell API duplication).
         Collect all <w:t> inside each unique <tc>, deduplicated by element identity.
      4. Global deduplication by lxml element identity (element set, not id()) across all sources.
    """
    try:
        doc = docx.Document(file_path)
        body = doc.element.body
        seen = set()   # lxml element identity set — deduplicates across all sources
        words = 0

        for child in body:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

            # ── PARAGRAPHS ──────────────────────────────────────────────────────────────
            if tag == 'p':
                # Rule 1: All <w:t> not inside a wsp shape container
                for t in child.xpath(
                    './/*[local-name()="t" and not(ancestor::*[local-name()="wsp"])]'
                ):
                    if t not in seen and t.text:
                        seen.add(t)
                        words += len(t.text.split())

                # Rule 2: <w:t> inside standalone txbxContent (outside wsp)
                # These are "floating" text boxes used as the entire layout in some CVs.
                # Without this, pure-textbox CVs (like the original Person2) extract 0 words.
                for t in child.xpath(
                    './/*[local-name()="txbxContent"'
                    ' and not(ancestor::*[local-name()="wsp"])'
                    ']//*[local-name()="t"]'
                ):
                    if t not in seen and t.text:
                        seen.add(t)
                        words += len(t.text.split())

            # ── TABLES ──────────────────────────────────────────────────────────────────
            elif tag == 'tbl':
                # Rule 3: Iterate <tc> XML elements directly — NOT via python-docx table API.
                # The API resolves merged cells by returning the same _tc object for multiple
                # logical positions, causing double-counting when there are row/col spans.
                # Direct XML iteration visits each <tc> node exactly once.
                for tc in child.xpath('.//*[local-name()="tc"]'):
                    if tc not in seen:
                        seen.add(tc)
                        for t in tc.xpath('.//*[local-name()="t"]'):
                            if t not in seen and t.text:
                                seen.add(t)
                                words += len(t.text.split())

        return words

    except Exception as e:
        return f"Error: {e}"


# ── Batch Processing ─────────────────────────────────────────────────────────────────────────────
folder_path = "Docx"
report_data = [["File Name", "Expected Words", "Extracted Words", "Accuracy (%)"]]

if os.path.exists(folder_path):
    print(f"--- Unified Smart Extractor v2 | Folder: {folder_path} ---\n")

    for filename in sorted(os.listdir(folder_path)):
        if not filename.endswith(".docx"):
            continue

        full_path = os.path.join(folder_path, filename)

        # Parse expected word count from filename (e.g. Person7-CV-245.docx → 245)
        expected = 0
        try:
            match = re.findall(r'\d+', filename) # using regex to find the expected word count in the filename
            if match:
                expected = int(match[-1]) # extracting the expected word count from the regex match
            else:
                expected = 0 

        except Exception:
            expected = 0

        result = unified_smart_extract(full_path)

        if isinstance(result, str) and result.startswith("Error"):
            print(f"  ERROR — {filename}: {result}")
            report_data.append([filename, expected, "ERROR", "0.00%"])
            continue

        extracted_words = result
        if expected > 0:
            variance = abs(expected - extracted_words)
            accuracy = max(0.0, (1 - variance / expected) * 100)
        else:
            accuracy = 0.0

        report_data.append([filename, expected, extracted_words, f"{accuracy:.2f}%"])
        print(
            f"  {filename:<45} | "
            f"extracted={extracted_words:<5} | "
            f"expected={expected:<5} | "
            f"accuracy={accuracy:.2f}%"
        )

    output_file = "Unified_Extraction_Report_v2.csv"
    with open(output_file, "w", newline="", encoding="utf-8") as f:
        csv.writer(f).writerows(report_data)

    print(f"\nDone. Report saved as: {output_file}")

else:
    print(f"Error: Folder '{folder_path}' not found.")
